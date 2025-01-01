import torch
import gc
import faiss
import os
import time
import pickle
import datasets
from tqdm import tqdm
from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig
from langchain_community.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores.utils import DistanceStrategy
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
from ragatouille import RAGPretrainedModel
from typing import Optional, List, Tuple

from PyPDF2 import PdfReader
import docx

# Constants
MARKDOWN_SEPARATORS = [
    "\n#{1,6} ",
    "```\n",
    "\n\\*\\*\\*+\n",
    "\n---+\n",
    "\n___+\n",
    "\n\n",
    "\n",
    " ",
    "",
]

EMBEDDING_MODEL_NAME = "thenlper/gte-small" #"sentence-transformers/all-MiniLM-L6-v2"
READER_MODEL_NAME = "meta-llama/Llama-3.2-3B-Instruct"
RERANKER_MODEL_NAME = "colbert-ir/colbertv2.0" #"cross-encoder/ms-marco-MiniLM-L-6-v2"
CACHE_PATH = r"C:\Users\Uvais\Documents\coding\streamlit\cache"
CHUNK_SIZE = 1024

READER_TOKENIZER = AutoTokenizer.from_pretrained(READER_MODEL_NAME)
RERANKER = RAGPretrainedModel.from_pretrained(RERANKER_MODEL_NAME)


def clear_memory():
    """Clear GPU, RAM, and VRAM before initializing models"""
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
        torch.cuda.synchronize()  # Ensure all tasks are finished
        print("GPU memory cleared!")

    gc.collect()
    print("RAM cleared!")


def load_process_data(data_dir: str) -> List[LangchainDocument]:
    """
    Loads personal data from PDFs, Word documents, and custom text files into LangchainDocument format.
    """
    documents = []

    for file in os.listdir(data_dir):
        file_path = os.path.join(data_dir, file)
        if file.endswith(".pdf"):
            reader = PdfReader(file_path)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    documents.append(LangchainDocument(page_content=text, metadata={"source": file}))

        elif file.endswith(".docx"):
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            if text:
                documents.append(LangchainDocument(page_content=text, metadata={"source": file}))

        elif file.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                if text:
                    documents.append(LangchainDocument(page_content=text, metadata={"source": file}))

    return documents


def split_documents(knowledge_base: List[LangchainDocument]) -> List[LangchainDocument]:
    """
    Splits the original documents into manageable text chunks.
    """
    emb_tokenizer = AutoTokenizer.from_pretrained(EMBEDDING_MODEL_NAME)
    text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
        emb_tokenizer,
        chunk_size=CHUNK_SIZE,
        chunk_overlap=int(CHUNK_SIZE / 8),
        add_start_index=True,
        strip_whitespace=True,
        separators=MARKDOWN_SEPARATORS,
    )
    docs_processed = []
    seen_texts = set()
    for doc in knowledge_base:
        chunks = text_splitter.split_documents([doc])
        for chunk in chunks:
            if chunk.page_content not in seen_texts:
                seen_texts.add(chunk.page_content)
                docs_processed.append(chunk)
    return docs_processed


def load_vector_database(cache_path: str) -> FAISS:
    """
    Loads an existing FAISS index plus associated docstore and mapping.
    """
    # Load the FAISS index
    index = faiss.read_index(os.path.join(cache_path, "vector_db.index"))

    # Load the docstore
    with open(os.path.join(cache_path, "docstore.pkl"), 'rb') as f:
        docstore = pickle.load(f)
        
    # Load the index-to-docstore mapping
    with open(os.path.join(cache_path, "index_to_docstore_id.pkl"), 'rb') as f:
        index_to_docstore_id = pickle.load(f)

    # Generate embedding model
    embedding_model = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL_NAME,
        multi_process=True,
        model_kwargs={"device": "cuda"},
        encode_kwargs={"normalize_embeddings": True},
    )

    vector_database_loaded = FAISS(
        index=index,
        embedding_function=embedding_model,
        docstore=docstore,
        index_to_docstore_id=index_to_docstore_id,
        distance_strategy= DistanceStrategy.DOT_PRODUCT #DistanceStrategy.COSINE
    )

    return vector_database_loaded


def generate_save_vector_database(data_dir: str) -> FAISS:
    """
    Processes documents, creates a FAISS index, and caches it locally.
    """
    # Load the data
    knowledge_base = load_process_data(data_dir)
    docs_processed = split_documents(knowledge_base)

    embedding_model = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL_NAME,
        multi_process=True,
        model_kwargs={"device": "cuda"},
        encode_kwargs={"normalize_embeddings": True},
    )

    vector_database = FAISS.from_documents(
        docs_processed,
        embedding_model,
        distance_strategy=DistanceStrategy.COSINE
    )
    
    # Cache the vector db
    faiss.write_index(vector_database.index, os.path.join(CACHE_PATH, "vector_db.index"))
    with open(os.path.join(CACHE_PATH, "docstore.pkl"), 'wb') as f:
        pickle.dump(vector_database.docstore, f)
    with open(os.path.join(CACHE_PATH, "index_to_docstore_id.pkl"), 'wb') as f:
        pickle.dump(vector_database.index_to_docstore_id, f)

    return vector_database


def initialise_llm():
    """
    Initializes and returns a text-generation pipeline with Llama-based model in 4-bit quantization.
    """
    bnb_config = BitsAndBytesConfig(
        load_in_4bit=True,
        bnb_4bit_use_double_quant=True,
        bnb_4bit_quant_type="nf4",
        bnb_4bit_compute_dtype=torch.float16,
    )

    model = AutoModelForCausalLM.from_pretrained(
        READER_MODEL_NAME,
        quantization_config=bnb_config,
        device_map="auto"
    )

    return pipeline(
        model=model,
        tokenizer=READER_TOKENIZER,
        task="text-generation",
        do_sample=True,
        temperature=0.3,
        repetition_penalty=1.1,
        return_full_text=False,
        max_new_tokens=300,
    )


def answer_with_rag(
    question: str, 
    llm, 
    knowledge_index: FAISS, 
    num_retrieved_docs: int = 15, 
    num_docs_final: int = 5,
    relevance_threshold: float = 0.7
) -> Tuple[str, List[LangchainDocument]]:
    """
    Retrieves relevant documents, optionally reranks them, then uses
    a LLM to generate an answer. The final prompt is built manually.
    """

    # Document Retrieval
    relevant_docs = knowledge_index.similarity_search(query=question, k=num_retrieved_docs)
    if not relevant_docs:
        return "I'm sorry, I cannot answer that based on the provided information.", []

    # Document Reranking
    if RERANKER:
        relevant_docs_strings = [doc.page_content for doc in relevant_docs]
        reranked_results = RERANKER.rerank(question, relevant_docs_strings, k=num_docs_final)

        # Sort by relevance score (descending) and apply threshold
        reranked_results = sorted(reranked_results, key=lambda x: x["score"], reverse=True)
        relevant_docs = [doc for doc in reranked_results if doc["score"] >= relevance_threshold]

    # Ensure the final set is within the required limit
    relevant_docs = relevant_docs[:num_docs_final]
    if not relevant_docs:
        return "No sufficiently relevant documents were found.", []

    context_text = ""
    for i, doc in enumerate(relevant_docs):
        context_text += f"Document {i}:\n{doc}\n"

    final_prompt = f"""System: You are an AI assistant called AI Butler, designed to showcase Uvais Karni's recent professional experiences, technical expertise, and academic achievements effectively to recruiters.
    - Focus on:
    - Highlighting Uvais Karni's **recent professional experiences** in AI, computer vision, machine learning, and data processing.
    - Prioritizing **recent technical accomplishments** over older roles or foundational qualifications.
    - Including academic achievements (e.g., Master's thesis) only when they directly relate to the query or add significant value.
    - Provide responses that align with job requirements and demonstrate how Uvais's experience and skills contribute to the role.
    - Use information from the context to craft concise, professional responses that emphasize Uvais's technical skills and contributions.
    - Only mention general qualifications like degrees if explicitly relevant to the question.
    - Do not make assumptions or fabricate details.
    - Format responses professionally:
    - Use **paragraphs** for detailed explanations of recent experiences, academic achievements, or technical concepts.
    - Use **bullet points** for key skills, tools, or methodologies.

    User:
    Context:
    {context_text}
    ---
    Recruiter's Question: {question}

    Assistant:
    """
    
    # LLM Answer Generation (simply pass the final_prompt to your pipeline)
    answer = llm(final_prompt)[0]["generated_text"]
    return answer, relevant_docs
