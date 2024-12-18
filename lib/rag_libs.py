import torch
import gc
import faiss
import os
import time
import pickle
import datasets
from tqdm import tqdm
from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig
from langchain_community.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores.utils import DistanceStrategy
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
from ragatouille import RAGPretrainedModel
from typing import Optional, List, Tuple

from PyPDF2 import PdfReader
import docx

# Constants
MARKDOWN_SEPARATORS = [
    "\n#{1,6} ",
    "```\n",
    "\n\\*\\*\\*+\n",
    "\n---+\n",
    "\n___+\n",
    "\n\n",
    "\n",
    " ",
    "",
]

EMBEDDING_MODEL_NAME = "thenlper/gte-small"
READER_MODEL_NAME = "HuggingFaceH4/zephyr-7b-beta"
CACHE_PATH = r"C:\Users\Uvais\Documents\coding\streamlit\cache"
CHUNK_SIZE = 512
READER_TOKENIZER = AutoTokenizer.from_pretrained(READER_MODEL_NAME)
RERANKER = RAGPretrainedModel.from_pretrained("colbert-ir/colbertv2.0")

# Define a more targeted prompt format for job interview questions
prompt_in_chat_format = [
    {
        "role": "system",
        "content": """You are an advanced AI system assisting Uvais Karni, a skilled developer with expertise. 
        - Your role is to present Uvais's qualifications effectively to recruiters for job opportunities. Tailor your responses to emphasize Uvais's **latest and most relevant experiences, skills, and projects**.
        - Focus on recent work and academic achievements unless the recruiter asks for older experiences.
        - Provide elaborate, clear, concise, and job-aligned responses.
        - Highlight Uvais’s contributions to open-source projects and his educational background in Computer Science from XYZ University.
        - Include specific details about his technical skills such as multi-object tracking, model optimization, and AI development.
        - Ensure your tone is formal and confident, demonstrating Uvais’s qualifications effectively.
        - Ensure the answer is well structure and contains bullets if necessary.
        - Do not responsed if you dont have informatiion for the question."""
    },
    {
        "role": "user",
        "content": """Context:
        {context}
        ---
        Recruiter's Question: {question}"""
    }
]

def clear_memory():
    """Clear GPU, RAM, and VRAM before initializing models"""
    # Clear GPU Memory (VRAM)
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
        torch.cuda.synchronize()  # Ensure all tasks are finished
        print("GPU memory cleared!")

    # Clear Python's RAM
    gc.collect()
    print("RAM cleared!")

def load_process_data(data_dir: str) -> List[LangchainDocument]:
    """
    Loads personal data from PDFs, Word documents, and custom text files into LangchainDocument format.

    Args:
        data_dir (str): Directory containing PDF, DOCX, and other personal text-based files.

    Returns:
        List[LangchainDocument]: List of documents for processing.
    """
    documents = []

    # Load PDF files
    for file in os.listdir(data_dir):
        file_path = os.path.join(data_dir, file)
        if file.endswith(".pdf"):
            reader = PdfReader(file_path)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    documents.append(LangchainDocument(page_content=text, metadata={"source": file}))

        # Load Word documents
        elif file.endswith(".docx"):
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            if text:
                documents.append(LangchainDocument(page_content=text, metadata={"source": file}))

        # Load custom text files (.txt)
        elif file.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                if text:
                    documents.append(LangchainDocument(page_content=text, metadata={"source": file}))

    return documents


def split_documents(
        knowledge_base: List[LangchainDocument]
        ) -> List[LangchainDocument]:
    
    emb_tokenizer = AutoTokenizer.from_pretrained(EMBEDDING_MODEL_NAME)
    text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
        emb_tokenizer,
        chunk_size=CHUNK_SIZE,
        chunk_overlap=int(CHUNK_SIZE / 10),
        add_start_index=True,
        strip_whitespace=True,
        separators=MARKDOWN_SEPARATORS,
    )
    docs_processed = []
    seen_texts = set()
    for doc in knowledge_base:
        chunks = text_splitter.split_documents([doc])
        for chunk in chunks:
            if chunk.page_content not in seen_texts:
                seen_texts.add(chunk.page_content)
                docs_processed.append(chunk)
    return docs_processed


def load_vector_database(cache_path):

    # Load the FAISS index
    index = faiss.read_index(os.path.join(cache_path,"vector_db.index"))

    # Load the docstore
    with open(os.path.join(cache_path,"docstore.pkl"), 'rb') as f:
        docstore = pickle.load(f)
        
    # Load the index-to-docstore mapping
    with open(os.path.join(cache_path,"index_to_docstore_id.pkl"), 'rb') as f:
        index_to_docstore_id = pickle.load(f)

    # generate embedding model
    embedding_model = HuggingFaceEmbeddings(
    model_name=EMBEDDING_MODEL_NAME,
    multi_process=True,
    model_kwargs={"device": "cuda"},
    encode_kwargs={"normalize_embeddings": True},
    )

    vector_database_loaded = FAISS(
    index=index,
    embedding_function=embedding_model,
    docstore=docstore,
    index_to_docstore_id=index_to_docstore_id,
    distance_strategy=DistanceStrategy.COSINE
    )

    return vector_database_loaded


def generate_save_vector_database(data_dir):
    
    # load the data
    knowledge_base = load_process_data(data_dir)
    # split tokenize and chunk
    docs_processed = split_documents(knowledge_base)

    embedding_model = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL_NAME,
        multi_process=True,
        model_kwargs={"device": "cuda"},
        encode_kwargs={"normalize_embeddings": True},
    )

    vector_database = FAISS.from_documents(docs_processed, embedding_model, distance_strategy=DistanceStrategy.COSINE)
    
    # cache the vector db
    # Save the FAISS index
    faiss.write_index(vector_database.index, os.path.join(CACHE_PATH,"vector_db.index"))
    # Save the docstore
    with open(os.path.join(CACHE_PATH,"docstore.pkl"), 'wb') as f:
        pickle.dump(vector_database.docstore, f)
    # Save the index-to-docstore mapping
    with open(os.path.join(CACHE_PATH,"index_to_docstore_id.pkl"), 'wb') as f:
        pickle.dump(vector_database.index_to_docstore_id, f)

    return vector_database


def initialise_llm():

    bnb_config = BitsAndBytesConfig(
        load_in_4bit=True,
        bnb_4bit_use_double_quant=True,
        bnb_4bit_quant_type="nf4",
        bnb_4bit_compute_dtype=torch.float16,
    )
    model = AutoModelForCausalLM.from_pretrained(READER_MODEL_NAME, quantization_config=bnb_config)
    return pipeline(
        model=model,
        tokenizer=READER_TOKENIZER,
        task="text-generation",
        do_sample=True,
        temperature=0.2,
        repetition_penalty=1.1,
        return_full_text=False,
        max_new_tokens=500,
    )


def answer_with_rag(
        question: str, 
        llm, 
        knowledge_index, 
        num_retrieved_docs=30, 
        num_docs_final=10) -> Tuple[str, List[LangchainDocument]]:
    
    # Document Retrieval
    relevant_docs = knowledge_index.similarity_search(query=question, k=num_retrieved_docs)

    # Document Reranking
    if RERANKER:
        relevant_docs_strings = [doc.page_content for doc in relevant_docs]
        relevant_docs = RERANKER.rerank(question, relevant_docs_strings, k=num_docs_final)

    relevant_docs = relevant_docs[:num_docs_final]

    # Build the final prompt
    context = "\nExtracted documents:\n" + "".join(
        [f"Document {i}:::\n{doc}" for i, doc in enumerate(relevant_docs)]
    )
    RAG_PROMPT_TEMPLATE = READER_TOKENIZER.apply_chat_template(
        prompt_in_chat_format, tokenize=False, add_generation_prompt=True
    )
    final_prompt = RAG_PROMPT_TEMPLATE.format(question=question, context=context)

    # LLM Answer Generation
    answer = llm(final_prompt)[0]["generated_text"]
    return answer, relevant_docs